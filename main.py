import os
import base64
import crewai
import streamlit as st
from io import BytesIO
from docx import Document
from dotenv import load_dotenv
from IPython.display import Markdown
from langchain_openai import ChatOpenAI
from crewai import Agent, Task, Crew
from crewai_tools import ScrapeWebsiteTool, SerperDevTool

# Load environment variables from .env file
load_dotenv()

# LLM object and API Key
os.environ['OPENAI_API_KEY'] = os.getenv('OPENAI_API_KEY')
os.environ['SERPER_API_KEY'] = os.getenv('SERPER_API_KEY')


def Generate_docx(content):
    doc = Document()
    # Add a title or heading if needed
    doc.add_heading('Diagnosis and Treatment Plan', level=1)

    # Add the main content
    doc.add_paragraph(content)

    # Save the document
    doc_file_path = "Diagnosis_and_Treatment_Plan.docx"
    doc.save(doc_file_path)

    return doc_file_path


# Download the docs
def get_download_link(bio, filename):
    b64 = base64.b64encode(bio.read()).decode()
    href = f'data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}'
    return href, filename


st.set_page_config(page_title="Healthcare Agent", layout="wide", page_icon="👨‍⚕️")

# TITLE
st.title("AI Agents to Empower Doctor's")

# TEXT INPUT
gender = st.selectbox("Select Gender", ("Male", "Female", "Other"))
age = st.number_input("Age", min_value=0, max_value=120, value=24)
symptoms = st.text_area("Enter Symptoms", placeholder="e.g., fever, cough, headache, stomach ache", label_visibility="visible")
medical_history = st.text_area("Enter Medical History", "e.g., diabetes, hypertension")

# Initialize Tools
search_tool = SerperDevTool()
scrape_tool = ScrapeWebsiteTool()

# Initialize LLM
llm = ChatOpenAI(
    model="gpt-3.5-turbo-16k",  # Specifies the model version used
    temperature=0.2,  # Controls randomness; 0 = deterministic, 1 = creative
    max_tokens=100,  # Maximum number of tokens (input + output)
    # top_p=0.9,  # Nucleus sampling; considers top X% probable next words
)

# Define Agents
diagnosticain = Agent(role="Medical Diagnostician",
                      goal="Analyze patient symptoms and medical history to provide a preliminary diagnosis.",
                      backstory="This Agent specializes in diagnosing medical conditions based on patient-reported symptoms and medical history. It uses advanced algorithms and medical knowledge to identify potential health issues.",
                      tools=[search_tool, scrape_tool],
                      allow_delegation=False,
                      llm=llm,
                      verbose=True)

treatment_advisor = Agent(role="Treatment Advisor",
                          goal="Recommend appropriate treatment plans based on the diagnosis provided by the Medical Diagnostician.",
                          backstory="This Agent specializes in creating treatment plans tailored to individual patient needs. It considers the diagnosis, patient history, and current best practices in medicine to recommend effective treatments.",
                          tools=[search_tool, scrape_tool],
                          allow_delegation=False,
                          llm=llm,
                          verbose=True)

# Define Tasks
diagnose_task = Task(
    description=(
        "1. Analyze the patient's symptoms ({symptoms}) and medical history ({medical_history}).\n"
        "2. Provide a preliminary diagnosis with possible conditions based on the provided information.\n"
        "3. Limit the diagnosis to the most likely conditions."
    ),
    expected_output="A preliminary diagnosis with a list of possible conditions.",
    agent=diagnosticain)

treatment_task = Task(
    description=(
        "1. Based on the diagnosis, recommend appropriate treatment plans step by step.\n"
        "2. Consider the patient's medical history ({medical_history}) and current symptoms ({symptoms}).\n"
        "3. Provide detailed treatment recommendations, including medications, lifestyle changes, and follow-up care."
    ),
    expected_output="A comprehensive treatment plan tailored to the patient's needs.",
    agent=treatment_advisor)

# Define Crew
crew = Crew(
    agents=[diagnosticain, treatment_advisor],
    tasks=[diagnose_task, treatment_task],
    verbose=True
)
# Execution Block


try:
    result = crew.kickoff(inputs={"symptoms": symptoms, "medical_history": medical_history})
    st.write("Result Type:", type(result))

    if isinstance(result, crewai.crews.crew_output.CrewOutput):
        st.write("CrewOutput received successfully!")
        st.write("Raw Output:", result.raw)

        combined_output = result.raw + "\n\n"

        if hasattr(result, 'tasks_output') and isinstance(result.tasks_output, list):
            for task_output in result.tasks_output:
                if isinstance(task_output, crewai.tasks.task_output.TaskOutput):
                    combined_output += task_output.raw + "\n\n"
                    st.subheader(f"Output from {task_output.agent}:")
                    st.markdown(task_output.raw, unsafe_allow_html=True)
                    st.write("---")

        docx_file = Generate_docx(combined_output)
        download_link, filename = get_download_link(open(docx_file, "rb"), "Diagnosis_and_Treatment_Plan.docx")
        st.markdown(download_link, unsafe_allow_html=True)
        st.success("Diagnosis and Treatment Plan has been generated and is ready for download.")
    else:
        st.error("Error: Unexpected result format. Please try again.")

except Exception as e:
    st.error(f"An error occurred: {str(e)}")
